import docx

#Document object cointains Paragraph objects
#Paragraph objects contain Run objects
d = docx.Document('c:\\users\\nicho\\Desktop\\PythonPrograms\\TestFiles\\demo.docx')

print(d.paragraphs)
#evaluates to list of paragraph objects
print(d.paragraphs[0])
print(d.paragraphs[0].text)

p = d.paragraphs[1]

print(p.runs)
#runs occur whenever the text changes at all (i.e. bold or italics)
print(p.runs[0])
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)

print(p.runs[1].bold)
#Returns True if that text is bolded
print(p.runs[0].bold)
#Returns None if text is not bold

#etc.
p.runs[0].italic
p.runs[3].underline

p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

print(p.style)
#Returns type of style (to see in word press ctrl+alt+shift+s)
p.style = 'Title'

d.save('c:\\users\\nicho\\Desktop\\PythonPrograms\\TestFiles\\demo2.docx')
