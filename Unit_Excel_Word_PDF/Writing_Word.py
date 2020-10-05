import docx
d = docx.Document()
#Creates blank new document
#Cannot insert paragraphs and run objects to places other than an end
#Need to Create new document, copy the paragraphs over, and insert
#during then

d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('c:\\users\\nicho\\Desktop\\PythonPrograms\\TestFiles\\demo3.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
#Now two run objects
p.runs[1].bold = True

d.save('c:\\users\\nicho\\Desktop\\PythonPrograms\\TestFiles\\demo3-1.docx')
